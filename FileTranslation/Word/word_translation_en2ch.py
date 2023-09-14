import sys
from docx import Document
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import my_utils
import transformers
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path)
    all_paragraphs = document.paragraphs
    paragraphs_text = []
    for paragraph in all_paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text.append([paragraph.style.name,paragraph.text])
    return paragraphs_text
if __name__ == '__main__':
    savePath = sys.argv[2]
    # 1、excel地址
    wordPath = sys.argv[1]
    # 提取文档信息
    text = get_paragraphs_text(wordPath)
    model_path = 'models/translate/en-zh'
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_path)
    translate_model = transformers.pipeline("translation", model=model, tokenizer=tokenizer)
    new_document = Document()
    for item in text:
        src_style = item[0]
        src_texts = [item[1]]
        print('标题层级：', src_style)
        print(src_texts)
        n_best = 1  # 每个输入样本的输出候选句子数量
        trg_texts = [my_utils.do_sentence(translate_model(src_texts[0])[0]['translation_text'])]
        print(trg_texts)
        if trg_texts == ['N']:
            trg_texts = ['']
        if src_style == 'Title':
            new_document.add_heading(trg_texts, 0)
        elif src_style[:7:] == 'Heading':
            new_document.add_heading(trg_texts, level=int(src_style[-1]))
        else:
            new_document.add_paragraph(trg_texts)
    word_name = os.path.basename(wordPath)
    new_document.save(os.path.join(savePath,word_name))