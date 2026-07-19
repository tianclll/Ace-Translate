# Copyright (c) 2025 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
import re
from collections import Counter
from pathlib import Path

from .base import BaseConverter, ConvertResult
from .math import convert_omath as _convert_omath
from .math import _M
from .registry import default_registry


# Regex patterns for Chinese numbered headings
_RE_H2 = re.compile(r"^[一二三四五六七八九十百千]+[、．.]")
_RE_H3 = re.compile(r"^（[一二三四五六七八九十百千]+）")

# Regex for field-code hyperlink instruction
_RE_FIELD_HYPERLINK = re.compile(r'HYPERLINK\s+"([^"]+)"')

# Regex for page-number-only footer/header text (e.g. "第  页", "共 页", "- 3 -", "Page of")
_RE_PAGE_ONLY = re.compile(r"^[\s第页共of\d\-/|·]*$", re.IGNORECASE)

# Word XML namespace
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = "{" + _W_NS + "}"

# Markup Compatibility namespace (mc:AlternateContent)
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_MC = "{" + _MC_NS + "}"

# WordprocessingShape namespace (wps:txbx)
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_WPS = "{" + _WPS_NS + "}"

# Chart namespace
_C_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
_C = "{" + _C_NS + "}"

# WordprocessingDrawing namespace (wp:inline, wp:anchor)
_WPD = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
# DrawingML main namespace
_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
# OPC relationships namespace
_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _escape_md_url(url: str) -> str:
    """Escape parentheses in URL for Markdown link syntax."""
    return url.replace("(", "%28").replace(")", "%29")


def _effective_bold(run, para) -> bool:
    """Resolve effective bold: run-level > character style > paragraph style."""
    b = run.bold
    if b is not None:
        return b
    try:
        if run.style and run.style.font.bold is not None:
            return run.style.font.bold
    except Exception:
        pass
    try:
        if para.style and para.style.font.bold is not None:
            return para.style.font.bold
    except Exception:
        pass
    return False


def _effective_italic(run, para) -> bool:
    """Resolve effective italic: run-level > character style > paragraph style."""
    i = run.italic
    if i is not None:
        return i
    try:
        if run.style and run.style.font.italic is not None:
            return run.style.font.italic
    except Exception:
        pass
    try:
        if para.style and para.style.font.italic is not None:
            return para.style.font.italic
    except Exception:
        pass
    return False


def _effective_underline(run, para) -> bool:
    """Resolve effective underline: run-level > character style > paragraph style."""
    u = run.underline
    if u is not None:
        return bool(u)
    try:
        if run.style and run.style.font.underline is not None:
            return bool(run.style.font.underline)
    except Exception:
        pass
    try:
        if para.style and para.style.font.underline is not None:
            return bool(para.style.font.underline)
    except Exception:
        pass
    return False


def _effective_superscript(run) -> bool:
    """Return True if run has superscript set (w:vertAlign w:val='superscript')."""
    try:
        return bool(run.font.superscript)
    except Exception:
        return False


def _effective_subscript(run) -> bool:
    """Return True if run has subscript set (w:vertAlign w:val='subscript')."""
    try:
        return bool(run.font.subscript)
    except Exception:
        return False


def _paragraph_has_math(para) -> bool:
    """Check if paragraph XML contains OMML math elements."""
    return para._element.find(f".//{_M}oMath") is not None


def _iter_math_paragraph_parts(para) -> list:
    """Parse a paragraph with mixed text/math content into a list of parts.

    Returns a list where each element is either:
    - ("text", items_list) — a group of text runs to be formatted
    - ("display_math", latex_str) — a display math block ($$...$$)
    - ("inline_math", latex_str) — an inline math expression ($...$)
    """
    from docx.text.run import Run
    from docx.text.hyperlink import Hyperlink

    text_items = []
    parts = []

    def flush_text():
        if text_items:
            parts.append(("text", list(text_items)))
            text_items.clear()

    for child in para._element:
        tag = child.tag
        local = tag.split("}")[-1] if "}" in tag else tag

        if tag == f"{_M}oMathPara":
            flush_text()
            for omath in child.findall(f"{_M}oMath"):
                latex = _convert_omath(omath)
                if latex:
                    parts.append(("display_math", latex))
        elif tag == f"{_M}oMath":
            flush_text()
            latex = _convert_omath(child)
            if latex:
                parts.append(("inline_math", latex))
        elif local == "r":
            try:
                run = Run(child, para)
                if run.text:
                    text_items.append(
                        (
                            _effective_bold(run, para),
                            _effective_italic(run, para),
                            _effective_underline(run, para),
                            bool(run.font.strike),
                            _effective_superscript(run),
                            _effective_subscript(run),
                            run.text,
                            "",
                        )
                    )
            except Exception:
                pass
        elif local == "hyperlink":
            try:
                hl = Hyperlink(child, para)
                try:
                    url = hl.url or ""
                except (KeyError, AttributeError):
                    url = ""
                for run in hl.runs:
                    if run.text:
                        text_items.append(
                            (
                                _effective_bold(run, para),
                                _effective_italic(run, para),
                                False,
                                bool(run.font.strike),
                                _effective_superscript(run),
                                _effective_subscript(run),
                                run.text,
                                url,
                            )
                        )
            except Exception:
                pass

    flush_text()
    return parts


def _paragraph_math_to_markdown(para) -> str:
    """Convert a paragraph containing OMML math to Markdown."""
    result = []
    for kind, data in _iter_math_paragraph_parts(para):
        if kind == "text":
            md = _runs_to_markdown(data)
            if md:
                result.append(md)
        elif kind == "display_math":
            result.append(f"$$\n{data}\n$$")
        elif kind == "inline_math":
            result.append(f"${data}$")
    return "".join(result)


def _paragraph_math_to_html(para) -> str:
    """Convert a paragraph containing OMML math to HTML inline text."""
    result = []
    for kind, data in _iter_math_paragraph_parts(para):
        if kind == "text":
            html = _runs_to_html(data)
            if html:
                result.append(html)
        elif kind == "display_math":
            result.append(f"$$\n{data}\n$$")
        elif kind == "inline_math":
            result.append(f"${data}$")
    return "".join(result)


def _get_body_font_size(doc) -> float:
    """Return the most common font size in the document (used as body size). Defaults to 16.0."""
    sizes: Counter = Counter()
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            if run.font.size:
                sizes[run.font.size.pt] += 1
                break  # only check the first run with an explicit size per paragraph
    return sizes.most_common(1)[0][0] if sizes else 16.0


def _detect_heading_level(para, body_font_size: float) -> int:
    """Return heading level (0 = not a heading, 1-6 = heading level)."""
    # Prefer Word built-in Heading styles
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    if style_name == "Title":
        return 1
    if style_name == "Subtitle":
        return 2

    text = para.text.strip()
    if not text:
        return 0

    # Use font size of the first run that has an explicit size
    font_size = None
    for run in para.runs:
        if run.font.size:
            font_size = run.font.size.pt
            break

    # Significantly larger than body font -> treat as heading (threshold: 1.5x, short paragraphs only)
    if font_size and font_size > body_font_size * 1.5:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 1
        if len(text) <= 60:
            return 2

    # Chinese numbered heading patterns
    if _RE_H2.match(text):
        return 2
    if _RE_H3.match(text):
        return 3

    return 0


class _FieldState:
    """Mutable state for tracking w:fldChar field codes across paragraphs."""

    __slots__ = ("active", "phase", "nest_depth", "url")
    active: bool
    phase: object  # None | "instr" | "result"
    nest_depth: int
    url: object  # None | str

    def __init__(self):
        self.active = False
        self.phase = None  # None | "instr" | "result"
        self.nest_depth = 0
        self.url = None


def _update_field_state_for_paragraph(para_element, field_state):
    """Update field_state by scanning a paragraph element's runs.

    Used for early-exit paragraphs (TOC / math / empty / code) to keep
    cross-paragraph field tracking accurate without building item lists.
    """
    for child in para_element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "r":
            continue
        fld_char = child.find(_W + "fldChar")
        if fld_char is not None:
            fld_type = fld_char.get(_W + "fldCharType")
            if fld_type == "begin":
                if field_state.phase == "result":
                    field_state.nest_depth += 1
                else:
                    field_state.active = True
                    field_state.phase = "instr"
                    field_state.url = None
            elif fld_type == "separate":
                if field_state.nest_depth == 0:
                    field_state.phase = "result"
            elif fld_type == "end":
                if field_state.nest_depth > 0:
                    field_state.nest_depth -= 1
                else:
                    field_state.active = False
                    field_state.phase = None
                    field_state.url = None
            continue
        instr_elem = child.find(_W + "instrText")
        if instr_elem is not None and field_state.phase == "instr":
            if instr_elem.text:
                m = _RE_FIELD_HYPERLINK.search(instr_elem.text)
                if m:
                    field_state.url = m.group(1)


def _iter_paragraph_items(para, field_state=None) -> list:
    """Extract (bold, italic, underline, strikethrough, superscript, subscript, text, url) tuples from a paragraph in document order.

    Handles python-docx Hyperlink objects and w:fldChar field-code hyperlinks.
    Silently degrades to plain text on error.
    Note: underline is forced to False inside Hyperlink/field-hyperlink runs to avoid Word's default hyperlink underline style.

    field_state: optional _FieldState instance for cross-paragraph field tracking.
                 If None, a fresh _FieldState is created (single-paragraph mode).
    """
    if field_state is None:
        field_state = _FieldState()

    def _split_breaks(items):
        """Expand items containing \\n (from <w:br/> soft line breaks) into per-line items with <br> separators."""
        expanded = []
        for (
            bold,
            italic,
            underline,
            strikethrough,
            superscript,
            subscript,
            text,
            url,
        ) in items:
            if "\n" not in text:
                expanded.append(
                    (
                        bold,
                        italic,
                        underline,
                        strikethrough,
                        superscript,
                        subscript,
                        text,
                        url,
                    )
                )
                continue
            segments = text.split("\n")
            for j, seg in enumerate(segments):
                if seg:
                    expanded.append(
                        (
                            bold,
                            italic,
                            underline,
                            strikethrough,
                            superscript,
                            subscript,
                            seg,
                            url,
                        )
                    )
                if j < len(segments) - 1:
                    expanded.append(
                        (False, False, False, False, False, False, "<br>\n", "")
                    )
        return expanded

    try:
        from docx.text.hyperlink import Hyperlink
    except ImportError:
        return _split_breaks(
            [
                (
                    _effective_bold(r, para),
                    _effective_italic(r, para),
                    _effective_underline(r, para),
                    bool(r.font.strike),
                    _effective_superscript(r),
                    _effective_subscript(r),
                    r.text,
                    "",
                )
                for r in para.runs
                if r.text
            ]
        )

    items = []
    try:
        content_iter = para.iter_inner_content()
    except Exception:
        return _split_breaks(
            [
                (
                    _effective_bold(r, para),
                    _effective_italic(r, para),
                    _effective_underline(r, para),
                    bool(r.font.strike),
                    _effective_superscript(r),
                    _effective_subscript(r),
                    r.text,
                    "",
                )
                for r in para.runs
                if r.text
            ]
        )

    for element in content_iter:
        try:
            if isinstance(element, Hyperlink):
                try:
                    url = element.url or ""
                except (KeyError, AttributeError):
                    url = ""
                for run in element.runs:
                    if not run.text:
                        continue
                    # Force underline=False: Word's Hyperlink style adds underline by default
                    items.append(
                        (
                            _effective_bold(run, para),
                            _effective_italic(run, para),
                            False,
                            bool(run.font.strike),
                            _effective_superscript(run),
                            _effective_subscript(run),
                            run.text,
                            url,
                        )
                    )
                # Fallback: hyperlink with no runs but has text
                if not element.runs and element.text:
                    items.append(
                        (False, False, False, False, False, False, element.text, url)
                    )
            else:
                # Plain Run — check for fldChar control elements first
                fld_char = element._element.find(_W + "fldChar")
                if fld_char is not None:
                    fld_type = fld_char.get(_W + "fldCharType")
                    if fld_type == "begin":
                        if field_state.phase == "result":
                            field_state.nest_depth += 1
                        else:
                            field_state.active = True
                            field_state.phase = "instr"
                            field_state.url = None
                    elif fld_type == "separate":
                        if field_state.nest_depth == 0:
                            field_state.phase = "result"
                    elif fld_type == "end":
                        if field_state.nest_depth > 0:
                            field_state.nest_depth -= 1
                        else:
                            field_state.active = False
                            field_state.phase = None
                            field_state.url = None
                    continue

                instr_elem = element._element.find(_W + "instrText")
                if instr_elem is not None:
                    if field_state.active and field_state.phase == "instr":
                        # Extract HYPERLINK url from instrText
                        if instr_elem.text:
                            m = _RE_FIELD_HYPERLINK.search(instr_elem.text)
                            if m:
                                field_state.url = m.group(1)
                    continue  # Never emit instrText run as content

                # Plain Run
                if not element.text:
                    continue

                # If we are in the result phase of a field-code hyperlink, apply URL
                # and suppress underline (same as w:hyperlink element handling above).
                if (
                    field_state.phase == "result"
                    and field_state.nest_depth == 0
                    and field_state.url
                ):
                    items.append(
                        (
                            _effective_bold(element, para),
                            _effective_italic(element, para),
                            False,  # suppress underline for field hyperlinks
                            bool(element.font.strike),
                            _effective_superscript(element),
                            _effective_subscript(element),
                            element.text,
                            field_state.url,
                        )
                    )
                else:
                    items.append(
                        (
                            _effective_bold(element, para),
                            _effective_italic(element, para),
                            _effective_underline(element, para),
                            bool(element.font.strike),
                            _effective_superscript(element),
                            _effective_subscript(element),
                            element.text,
                            "",
                        )
                    )
        except Exception:
            continue

    return _split_breaks(items)


def _merge_runs(items) -> list:
    """Merge adjacent items with identical (bold, italic, underline, strikethrough, superscript, subscript, url).

    Returns [(bold, italic, underline, strikethrough, superscript, subscript, text, url)].
    """
    merged: list[tuple[bool, bool, bool, bool, bool, bool, str, str]] = []
    for (
        bold,
        italic,
        underline,
        strikethrough,
        superscript,
        subscript,
        text,
        url,
    ) in items:
        if not text:
            continue
        if (
            merged
            and merged[-1][0] == bold
            and merged[-1][1] == italic
            and merged[-1][2] == underline
            and merged[-1][3] == strikethrough
            and merged[-1][4] == superscript
            and merged[-1][5] == subscript
            and merged[-1][7] == url
        ):
            merged[-1] = (
                bold,
                italic,
                underline,
                strikethrough,
                superscript,
                subscript,
                merged[-1][6] + text,
                url,
            )
        else:
            merged.append(
                (
                    bold,
                    italic,
                    underline,
                    strikethrough,
                    superscript,
                    subscript,
                    text,
                    url,
                )
            )
    return merged


def _runs_to_markdown(items) -> str:
    """Convert paragraph items to Markdown inline text, merging adjacent items with identical formatting."""
    parts = []
    for (
        bold,
        italic,
        underline,
        strikethrough,
        superscript,
        subscript,
        text,
        url,
    ) in _merge_runs(items):
        if bold or italic or underline or strikethrough or superscript or subscript:
            # CommonMark: marker characters must not be surrounded by spaces
            leading = len(text) - len(text.lstrip())
            trailing = len(text) - len(text.rstrip())
            prefix = text[:leading] if leading else ""
            suffix = text[len(text) - trailing :] if trailing else ""
            inner = text.strip()
            if inner:
                # Apply strikethrough first (innermost)
                if strikethrough:
                    inner = f"~~{inner}~~"
                # Apply bold/italic
                if bold and italic:
                    inner = f"***{inner}***"
                elif bold:
                    inner = f"**{inner}**"
                elif italic:
                    inner = f"*{inner}*"
                # Apply underline
                if underline:
                    inner = f"<u>{inner}</u>"
                # Apply superscript/subscript (outermost)
                if superscript:
                    inner = f"<sup>{inner}</sup>"
                elif subscript:
                    inner = f"<sub>{inner}</sub>"
                text = prefix + inner + suffix
            elif underline and text:
                # Pure whitespace + underline = fill-in line (e.g. "作者姓名：___")
                # Replace spaces with NBSP so Markdown renderers preserve width
                text = "<u>" + "\u00a0" * len(text) + "</u>"
        if url:
            text = f"[{text}]({_escape_md_url(url)})"
        parts.append(text)
    # Prevent bold/italic/strikethrough markers from merging with adjacent alphanumeric text (CommonMark requirement)
    result = []
    for i, part in enumerate(parts):
        if i > 0 and result:
            prev = result[-1]
            # Previous part ends with closing marker and current part starts with alphanumeric
            if prev.endswith(("**", "*", "~~")) and part and part[0].isalnum():
                result.append("\u200b")
        result.append(part)
    return "".join(result)


def _runs_to_html(items) -> str:
    """Convert paragraph items to HTML inline text."""
    parts = []
    for (
        bold,
        italic,
        underline,
        strikethrough,
        superscript,
        subscript,
        text,
        url,
    ) in _merge_runs(items):
        if bold:
            text = f"<b>{text}</b>"
        if italic:
            text = f"<i>{text}</i>"
        if underline:
            text = f"<u>{text}</u>"
        if strikethrough:
            text = f"<del>{text}</del>"
        if superscript:
            text = f"<sup>{text}</sup>"
        if subscript:
            text = f"<sub>{text}</sub>"
        if url:
            text = f'<a href="{url}">{text}</a>'
        parts.append(text)
    return "".join(parts)


# Regex for TOC style names (e.g. "toc 1", "TOC 2", "TOC3")
_RE_TOC_STYLE = re.compile(r"(?i)^toc\s*(\d+)$")

# Flat TOC styles with no level concept (e.g. figure/table of contents)
_TOC_FLAT_STYLES = {"table of figures"}

# Regex for PAGEREF anchor in field instructions
_RE_PAGEREF = re.compile(r"PAGEREF\s+(_Toc\w+)")


def _is_toc_paragraph(para):
    """Return TOC level (1-9) if paragraph uses a TOC style, else None."""
    style_name = (para.style.name if para.style else "").strip()
    m = _RE_TOC_STYLE.match(style_name)
    if m:
        return int(m.group(1))
    if style_name.lower() in _TOC_FLAT_STYLES:
        return 1
    return None


def _extract_toc_text(para) -> str:
    """Extract display text from a TOC paragraph, stripping trailing page numbers."""
    text = para.text
    # Remove trailing page number: split on last tab, discard if it's a pure number
    if "\t" in text:
        before_tab, _, after_tab = text.rpartition("\t")
        if after_tab.strip().isdigit():
            text = before_tab
    return text.strip()


def _extract_toc_anchor(para_element):
    """Extract anchor name from a TOC paragraph element.

    Priority:
    1. w:hyperlink[@w:anchor] attribute
    2. PAGEREF _TocXXXX in w:instrText
    """
    # Check for w:hyperlink with anchor attribute
    for hl in para_element.findall(f".//{_W}hyperlink"):
        anchor = hl.get(f"{_W}anchor")
        if anchor and anchor.startswith("_Toc"):
            return anchor

    # Fall back to PAGEREF field instruction
    for instr in para_element.findall(f".//{_W}instrText"):
        if instr.text:
            m = _RE_PAGEREF.search(instr.text)
            if m:
                return m.group(1)
    return None


def _toc_entries_to_markdown(entries: list) -> str:
    """Convert list of (text, anchor, level) TOC entries to Markdown list."""
    lines = []
    for text, anchor, level in entries:
        indent = "  " * (level - 1)
        if anchor:
            lines.append(f"{indent}- [{text}](#{anchor})")
        else:
            lines.append(f"{indent}- {text}")
    return "\n".join(lines)


def _extract_heading_toc_bookmarks(para_element):
    """Return all _Toc bookmark names found in a paragraph element.

    A heading may carry multiple _Toc bookmarks from repeated TOC updates
    (each update inserts a new bookmark without removing old ones).  Returning
    all of them ensures that TOC entries from any generation can link here.
    """
    names = []
    for bm in para_element.findall(f".//{_W}bookmarkStart"):
        name = bm.get(f"{_W}name", "")
        if name.startswith("_Toc"):
            names.append(name)
    return names


def _flatten_body(body):
    """Yield body children, expanding sdt elements into their sdtContent children."""
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "sdt":
            sdt_content = child.find(f"{_W}sdtContent")
            if sdt_content is not None:
                yield from _flatten_body(sdt_content)
        else:
            yield child


_CODE_FONTS = {
    "Courier New",
    "Courier",
    "Consolas",
    "Monaco",
    "Menlo",
    "Source Code Pro",
    "Fira Code",
    "DejaVu Sans Mono",
    "monospace",
}


def _is_code_paragraph(para) -> bool:
    """Return True if all text-bearing runs in the paragraph use a monospace font."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    runs_with_font = [r for r in runs_with_text if r.font.name]
    # At least one run must have an explicit font, and all such runs must be monospace
    if not runs_with_font:
        return False
    return all(r.font.name in _CODE_FONTS for r in runs_with_font)


def _get_content_width(doc) -> int:
    """Return the content area width of the document in EMU."""
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _table_to_html(
    table, doc, image_counter: list, images: dict, content_width: int = 0
) -> str:
    """Convert a python-docx Table to an HTML table, handling merged cells and inline images."""
    grid = [[cell for cell in row.cells] for row in table.rows]
    nrows = len(grid)
    if nrows == 0:
        return ""
    ncols = len(grid[0])

    visited: set[tuple[int, int]] = set()
    html_parts = ["<table>"]

    for i in range(nrows):
        html_parts.append("<tr>")
        for j in range(ncols):
            if (i, j) in visited:
                continue
            tc = grid[i][j]._tc
            # Compute colspan
            colspan = 1
            while j + colspan < ncols and grid[i][j + colspan]._tc is tc:
                visited.add((i, j + colspan))
                colspan += 1
            # Compute rowspan
            rowspan = 1
            while i + rowspan < nrows and grid[i + rowspan][j]._tc is tc:
                for k in range(colspan):
                    visited.add((i + rowspan, j + k))
                rowspan += 1

            cell = grid[i][j]
            content_parts = []
            for para in cell.paragraphs:
                img_list = _extract_images_from_paragraph(para, doc, image_counter)
                for filename, img_bytes, cx_emu in img_list:
                    rel_path = f"assets/{filename}"
                    images[rel_path] = img_bytes
                    if cx_emu and content_width:
                        pct = min(round(cx_emu / content_width * 100), 100)
                        content_parts.append(
                            f'<img src="assets/{filename}" width="{pct}%">'
                        )
                    else:
                        content_parts.append(f'<img src="assets/{filename}">')
                para_html = (
                    _paragraph_math_to_html(para).strip()
                    if _paragraph_has_math(para)
                    else _runs_to_html(_iter_paragraph_items(para)).strip()
                    or para.text.strip()
                )
                if para_html:
                    content_parts.append(para_html)
                # Extract text box content from table cell paragraph
                try:
                    tb_groups = _extract_textbox_paragraphs(para._element)
                    for tb_paras in tb_groups:
                        for tb_p in tb_paras:
                            tb_t_elems = tb_p.findall(f".//{_W}t")
                            tb_text = "".join(
                                t.text for t in tb_t_elems if t.text
                            ).strip()
                            if tb_text:
                                content_parts.append(f"[{tb_text}]")
                except Exception:
                    pass
            cell_html = "<br>".join(content_parts) if content_parts else ""

            is_header = False
            trPr = table.rows[i]._tr.find(f"{_W}trPr")
            if trPr is not None:
                tbl_header = trPr.find(f"{_W}tblHeader")
                if tbl_header is not None:
                    val = tbl_header.get(f"{_W}val")
                    is_header = val is None or val.lower() not in ("false", "0", "off")
            if not is_header and i == 0 and nrows > 1:
                is_header = (
                    True  # fallback: treat first row as header in multi-row tables
                )
            tag = "th" if is_header else "td"
            attrs = ""
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            html_parts.append(f"<{tag}{attrs}>{cell_html}</{tag}>")
        html_parts.append("</tr>")

    html_parts.append("</table>")
    return "\n".join(html_parts)


_MIME_TO_EXT = {
    "image/jpeg": "jpeg",
    "image/jpg": "jpg",
    "image/png": "png",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
    "image/webp": "webp",
    "image/svg+xml": "svg",
    "image/x-emf": "emf",
    "image/x-wmf": "wmf",
}


def _extract_images_from_paragraph(para, doc, image_counter: list) -> list:
    """Extract images from a paragraph. Returns [(filename, bytes, cx_emu)] where cx_emu is 0 if unknown."""
    results = []
    containers = para._element.findall(f".//{_WPD}inline") + para._element.findall(
        f".//{_WPD}anchor"
    )
    for container in containers:
        extent = container.find(f"{_WPD}extent")
        cx_emu = int(extent.get("cx", 0)) if extent is not None else 0

        blip = container.find(f".//{_A}blip")
        if blip is None:
            continue
        r_embed = blip.get(f"{_R}embed")
        if not r_embed:
            continue
        try:
            rel = doc.part.rels[r_embed]
            img_bytes = rel.target_part.blob
            content_type = rel.target_part.content_type
            ext = _MIME_TO_EXT.get(content_type, "png")
            image_counter[0] += 1
            filename = f"image{image_counter[0]}.{ext}"
            results.append((filename, img_bytes, cx_emu))
        except (KeyError, AttributeError):
            pass
    return results


def _extract_chart_tables(para, doc) -> list:
    """Extract chart data as HTML tables from drawings in a paragraph."""
    try:
        from datetime import datetime, timedelta

        results = []
        containers = para._element.findall(f".//{_WPD}inline") + para._element.findall(
            f".//{_WPD}anchor"
        )

        for container in containers:
            chart_ref = container.find(f".//{_C}chart")
            if chart_ref is None:
                continue
            r_id = chart_ref.get(f"{_R}id")
            if not r_id:
                continue
            try:
                rel = doc.part.rels[r_id]
                chart_part = rel.target_part
            except (KeyError, AttributeError):
                continue

            try:
                import lxml.etree as etree

                chart_root = etree.fromstring(chart_part.blob)
            except Exception:
                continue

            # Extract chart title
            title_text = ""
            title_el = chart_root.find(f".//{_C}title")
            if title_el is not None:
                a_t_els = title_el.findall(f".//{_A}t")
                title_text = "".join((el.text or "") for el in a_t_els).strip()

            # Extract axis titles
            cat_ax_title = ""
            for ax_tag in (f"{_C}catAx", f"{_C}dateAx"):
                ax_el = chart_root.find(f".//{ax_tag}")
                if ax_el is not None:
                    t_el = ax_el.find(f"{_C}title")
                    if t_el is not None:
                        texts = t_el.findall(f".//{_A}t")
                        cat_ax_title = "".join(el.text or "" for el in texts).strip()
                        break

            val_ax_title = ""
            val_ax_el = chart_root.find(f".//{_C}valAx")
            if val_ax_el is not None:
                t_el = val_ax_el.find(f"{_C}title")
                if t_el is not None:
                    texts = t_el.findall(f".//{_A}t")
                    val_ax_title = "".join(el.text or "" for el in texts).strip()

            # Extract series data
            series_list = chart_root.findall(f".//{_C}ser")
            if not series_list:
                continue

            # Collect categories from first series
            categories = []
            cat_is_date = False
            date_format_code = ""
            first_ser = series_list[0]
            cat_el = first_ser.find(f"{_C}cat")
            if cat_el is not None:
                str_cache = cat_el.find(f".//{_C}strCache")
                num_cache = cat_el.find(f".//{_C}numCache")
                if str_cache is not None:
                    for pt in str_cache.findall(f"{_C}pt"):
                        v = pt.find(f"{_C}v")
                        categories.append(v.text if v is not None else "")
                elif num_cache is not None:
                    fmt_el = num_cache.find(f"{_C}formatCode")
                    if fmt_el is not None:
                        fc = fmt_el.text or ""
                        date_format_code = fc
                        cat_is_date = any(k in fc.lower() for k in ("y", "m", "d"))
                    for pt in num_cache.findall(f"{_C}pt"):
                        v = pt.find(f"{_C}v")
                        if v is not None and v.text:
                            try:
                                serial = float(v.text)
                                if cat_is_date:
                                    dt = datetime(1899, 12, 30) + timedelta(days=serial)
                                    categories.append(dt.strftime("%Y-%m-%d"))
                                else:
                                    categories.append(v.text)
                            except (ValueError, OverflowError):
                                categories.append(v.text or "")
                        else:
                            categories.append("")

            # Collect series names and values
            series_names = []
            series_values = []
            for ser in series_list:
                # Series name: val axis title > <c:tx> > ""
                name = val_ax_title
                if not name:
                    tx_el = ser.find(f"{_C}tx")
                    if tx_el is not None:
                        str_ref = tx_el.find(f".//{_C}strCache")
                        if str_ref is not None:
                            pt = str_ref.find(f"{_C}pt")
                            if pt is not None:
                                v = pt.find(f"{_C}v")
                                name = v.text if v is not None else ""
                        else:
                            v = tx_el.find(f"{_C}v")
                            if v is not None:
                                name = v.text or ""
                series_names.append(name)

                # Values
                vals = []
                val_el = ser.find(f"{_C}val")
                if val_el is not None:
                    num_cache = val_el.find(f".//{_C}numCache")
                    if num_cache is not None:
                        pts = {
                            int(pt.get("idx", 0)): pt
                            for pt in num_cache.findall(f"{_C}pt")
                        }
                        max_idx = max(pts.keys()) if pts else -1
                        for idx in range(max_idx + 1):
                            pt = pts.get(idx)
                            if pt is not None:
                                v = pt.find(f"{_C}v")
                                vals.append(v.text if v is not None else "")
                            else:
                                vals.append("")
                series_values.append(vals)

            if not series_names:
                continue

            # Build HTML table
            n_rows = max(
                len(categories),
                max((len(v) for v in series_values), default=0),
            )
            html_parts = ["<table>"]
            if title_text:
                html_parts.append(f"<caption>{title_text}</caption>")
            # Header row — omit <thead> entirely if no header info
            has_header = cat_ax_title or any(name for name in series_names)
            if has_header:
                html_parts.append("<thead><tr>")
                html_parts.append(f"<th>{cat_ax_title}</th>")
                for name in series_names:
                    html_parts.append(f"<th>{name}</th>")
                html_parts.append("</tr></thead>")
            # Data rows
            html_parts.append("<tbody>")
            for i in range(n_rows):
                cat = categories[i] if i < len(categories) else ""
                html_parts.append(f"<tr><td>{cat}</td>")
                for vals in series_values:
                    val = vals[i] if i < len(vals) else ""
                    html_parts.append(f"<td>{val}</td>")
                html_parts.append("</tr>")
            html_parts.append("</tbody></table>")
            results.append("".join(html_parts))

        return results
    except Exception:
        return []


def _extract_textbox_paragraphs(element) -> list:
    """Extract paragraphs from text boxes embedded in a body element.

    Text boxes appear as mc:AlternateContent > mc:Choice > wps:txbx > w:txbxContent > w:p.
    Only mc:Choice is used to avoid duplicating VML fallback content.

    Returns list of lists of w:p elements, one list per non-empty text box.
    """
    try:
        result = []
        for alt_content in element.iter(f"{_MC}AlternateContent"):
            choice = alt_content.find(f"{_MC}Choice")
            if choice is None:
                continue
            for txbx in choice.iter(f"{_WPS}txbx"):
                txbx_content = txbx.find(f"{_W}txbxContent")
                if txbx_content is None:
                    continue
                paras = txbx_content.findall(f"{_W}p")
                if not paras:
                    continue
                # Skip boxes where all paragraphs are empty
                has_text = any(
                    p.find(f".//{_W}t") is not None
                    and any(t.text for t in p.findall(f".//{_W}t") if t.text)
                    for p in paras
                )
                if has_text:
                    result.append(paras)
        return result
    except Exception:
        return []


def _textbox_paragraphs_to_markdown(
    textbox_groups, doc, body_font_size, image_counter, images, content_width
) -> list:
    """Convert text box paragraph groups to blockquote Markdown lines.

    Returns a list of strings (lines) to append to the main output.
    """
    try:
        from docx.text.paragraph import Paragraph
    except ImportError:
        return []

    output_lines = []
    for group_idx, para_elements in enumerate(textbox_groups):
        group_lines = []
        for p_elem in para_elements:
            try:
                para = Paragraph(p_elem, doc)

                # Handle images in text box
                img_list = _extract_images_from_paragraph(para, doc, image_counter)
                for filename, img_bytes, cx_emu in img_list:
                    rel_path = f"assets/{filename}"
                    images[rel_path] = img_bytes
                    if cx_emu and content_width:
                        pct = min(round(cx_emu / content_width * 100), 100)
                        group_lines.append(
                            f'> <img src="assets/{filename}" width="{pct}%">'
                        )
                    else:
                        group_lines.append(f'> <img src="assets/{filename}">')

                # Handle math formulas
                if _paragraph_has_math(para):
                    math_md = _paragraph_math_to_markdown(para)
                    if math_md:
                        group_lines.append(f"> {math_md}")
                    continue

                # Plain inline text (no heading/list/code detection for text boxes)
                inline = _runs_to_markdown(_iter_paragraph_items(para))
                if not inline:
                    inline = para.text.strip()
                if inline:
                    group_lines.append(f"> {inline}")
            except Exception:
                continue

        if group_lines:
            if output_lines:
                output_lines.append("")
            output_lines.extend(group_lines)

    return output_lines


def _build_numbering_map(doc) -> dict:
    """Parse numbering.xml and return {numId: {ilvl: numFmt}} mapping."""
    numbering_map = {}
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return numbering_map
    numbering_elem = numbering_part._element

    # Build abstractNumId -> {ilvl: numFmt}
    abstract = {}
    for abs_num in numbering_elem.findall(f"{_W}abstractNum"):
        abs_id = abs_num.get(f"{_W}abstractNumId")
        levels = {}
        for lvl in abs_num.findall(f"{_W}lvl"):
            ilvl = int(lvl.get(f"{_W}ilvl", "0"))
            fmt_elem = lvl.find(f"{_W}numFmt")
            fmt = (
                fmt_elem.get(f"{_W}val", "bullet") if fmt_elem is not None else "bullet"
            )
            levels[ilvl] = fmt
        abstract[abs_id] = levels

    # Map numId -> abstractNumId
    for num in numbering_elem.findall(f"{_W}num"):
        num_id = num.get(f"{_W}numId")
        abs_ref = num.find(f"{_W}abstractNumId")
        if abs_ref is not None:
            abs_id = abs_ref.get(f"{_W}val")
            if abs_id in abstract:
                numbering_map[num_id] = abstract[abs_id]
    return numbering_map


def _get_list_info(para, numbering_map) -> tuple | None:
    """Return (list_type, ilvl, num_id) or None. list_type: 'bullet' | 'ordered'"""
    pPr = para._element.find(f"{_W}pPr")
    if pPr is None:
        return None
    numPr = pPr.find(f"{_W}numPr")
    if numPr is None:
        return None
    numId_elem = numPr.find(f"{_W}numId")
    ilvl_elem = numPr.find(f"{_W}ilvl")
    if numId_elem is None:
        return None
    num_id = numId_elem.get(f"{_W}val")
    ilvl = int(ilvl_elem.get(f"{_W}val", "0")) if ilvl_elem is not None else 0
    if num_id not in numbering_map:
        return None
    fmt = numbering_map[num_id].get(ilvl, "bullet")
    list_type = (
        "ordered"
        if fmt in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman")
        else "bullet"
    )
    return (list_type, ilvl, num_id)


def _convert_body(doc, *, extract_drawings=True) -> tuple:
    """Traverse body elements in order and produce Markdown. Returns (markdown_str, images_dict)."""
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        raise RuntimeError(
            "DOCX conversion requires python-docx: pip install paddleocr[doc2md]"
        )

    body_font_size = _get_body_font_size(doc)
    content_width = _get_content_width(doc)
    numbering_map = _build_numbering_map(doc)
    lines: list[str] = []
    images: dict = {}
    image_counter = [0]  # wrapped in list so inner functions can mutate it
    code_buf: list[str] = []  # buffer for consecutive code paragraphs
    toc_buf: list[tuple] = []  # buffer for consecutive TOC paragraphs
    ol_counters: dict[str, int] = {}  # key = "{numId}-{ilvl}", value = current index
    prev_was_list = False

    def flush_code_buf():
        """Flush the code buffer as a fenced code block."""
        if code_buf:
            lines.append("```")
            lines.extend(code_buf)
            lines.append("```")
            lines.append("")
            code_buf.clear()

    def flush_toc_buf():
        """Flush the TOC buffer as a Markdown list."""
        if toc_buf:
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(_toc_entries_to_markdown(toc_buf))
            lines.append("")
            toc_buf.clear()

    field_state = _FieldState()
    for child in _flatten_body(doc.element.body):
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = Paragraph(child, doc)

            # TOC paragraph detection — before image extraction and other processing
            toc_level = _is_toc_paragraph(para)
            if toc_level is not None:
                toc_text = _extract_toc_text(para)
                if toc_text:
                    toc_anchor = _extract_toc_anchor(child)
                    toc_buf.append((toc_text, toc_anchor, toc_level))
                _update_field_state_for_paragraph(child, field_state)
                continue

            # Non-TOC paragraph: flush any buffered TOC entries
            flush_toc_buf()

            # Extract text box content for this paragraph element
            pending_textbox_lines = []
            if extract_drawings:
                tb_groups = _extract_textbox_paragraphs(child)
                if tb_groups:
                    pending_textbox_lines = _textbox_paragraphs_to_markdown(
                        tb_groups,
                        doc,
                        body_font_size,
                        image_counter,
                        images,
                        content_width,
                    )

            def _flush_textbox():
                if pending_textbox_lines:
                    if lines and lines[-1] != "":
                        lines.append("")
                    lines.extend(pending_textbox_lines)
                    lines.append("")

            # Extract images first
            img_list = _extract_images_from_paragraph(para, doc, image_counter)
            for filename, img_bytes, cx_emu in img_list:
                flush_code_buf()
                rel_path = f"assets/{filename}"
                images[rel_path] = img_bytes
                if cx_emu and content_width:
                    pct = min(round(cx_emu / content_width * 100), 100)
                    lines.append(f'<img src="assets/{filename}" width="{pct}%">')
                else:
                    lines.append(f'<img src="assets/{filename}">')
                lines.append("")

            # Extract chart data tables
            chart_tables = _extract_chart_tables(para, doc)
            for chart_html in chart_tables:
                flush_code_buf()
                lines.append(chart_html)
                lines.append("")

            text = para.text.strip()

            # Math formula detection — must check before skipping empty-text paragraphs
            # (pure formula paragraphs have para.text == "")
            if _paragraph_has_math(para):
                flush_code_buf()
                if prev_was_list:
                    lines.append("")
                prev_was_list = False
                math_md = _paragraph_math_to_markdown(para)
                if math_md:
                    lines.append(math_md)
                    lines.append("")
                _flush_textbox()
                _update_field_state_for_paragraph(child, field_state)
                continue

            if not text:
                # Emit any _Toc bookmark anchors even for empty paragraphs.
                # Headings whose visible text comes from list numbering (not w:t
                # runs) have para.text == "" but still carry _Toc bookmarks that
                # TOC entries link to.
                for toc_bm in _extract_heading_toc_bookmarks(child):
                    lines.append(f'<a id="{toc_bm}"></a>')
                _flush_textbox()
                if not img_list and not pending_textbox_lines:
                    if code_buf:
                        code_buf.append("")  # preserve blank lines inside code blocks
                    elif lines and lines[-1] != "":
                        lines.append("")
                _update_field_state_for_paragraph(child, field_state)
                continue

            # Code paragraph: buffer it without heading/inline formatting
            if _is_code_paragraph(para):
                code_buf.append(para.text)
                _flush_textbox()
                _update_field_state_for_paragraph(child, field_state)
                continue

            # Non-code paragraph: flush any buffered code first
            flush_code_buf()

            level = _detect_heading_level(para, body_font_size)
            inline = _runs_to_markdown(_iter_paragraph_items(para, field_state)) or text

            if level > 0:
                # Strip outer **...** wrapping that headings may have inherited
                clean = inline.strip()
                if clean.startswith("**") and clean.endswith("**"):
                    clean = clean[2:-2]
                # Heading lines cannot span multiple source lines; revert <br>\n → <br>
                clean = clean.replace("<br>\n", "<br>")
                if prev_was_list:
                    lines.append("")
                prev_was_list = False
                # Add _Toc bookmark anchors if present (makes TOC links jumpable).
                # A heading may have multiple _Toc bookmarks from repeated TOC updates.
                for toc_bm in _extract_heading_toc_bookmarks(child):
                    lines.append(f'<a id="{toc_bm}"></a>')
                lines.append(f"{'#' * level} {clean}")
                lines.append("")
            else:
                list_info = _get_list_info(para, numbering_map)
                if list_info:
                    list_type, ilvl, num_id = list_info
                    indent = "    " * ilvl
                    if list_type == "ordered":
                        counter_key = f"{num_id}-{ilvl}"
                        ol_counters[counter_key] = ol_counters.get(counter_key, 0) + 1
                        prefix = f"{indent}{ol_counters[counter_key]}. "
                    else:
                        prefix = f"{indent}- "
                    if not prev_was_list and lines and lines[-1] != "":
                        lines.append("")
                    prev_was_list = True
                    # List item continuation lines need indentation; revert <br>\n → <br>
                    lines.append(prefix + inline.replace("<br>\n", "<br>"))
                else:
                    if prev_was_list:
                        lines.append("")
                    prev_was_list = False
                    # Reset ordered list counters when a list is interrupted
                    ol_counters.clear()
                    # Add _Toc bookmark anchors for non-heading paragraphs (e.g. Caption)
                    for toc_bm in _extract_heading_toc_bookmarks(child):
                        lines.append(f'<a id="{toc_bm}"></a>')
                    lines.append(inline)
                    lines.append("")

            _flush_textbox()

        elif tag == "tbl":
            flush_code_buf()
            flush_toc_buf()
            if prev_was_list:
                lines.append("")
            prev_was_list = False
            ol_counters.clear()
            table = Table(child, doc)
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(
                _table_to_html(table, doc, image_counter, images, content_width)
            )
            lines.append("")

    # Strip trailing blank lines
    flush_code_buf()
    flush_toc_buf()
    while lines and lines[-1] == "":
        lines.pop()

    return "\n".join(lines), images


def _extract_headers_footers(doc) -> tuple:
    """Extract header and footer text from all document sections.

    Returns (header_lines: list[str], footer_lines: list[str]).
    Deduplicates content and filters out page-number-only text.
    """

    def _collect_from_parts(parts, seen: set) -> list:
        results = []
        for part in parts:
            try:
                if part.is_linked_to_previous:
                    continue
                texts = []
                for para in part.paragraphs:
                    try:
                        items = list(_iter_paragraph_items(para))
                        inline = _runs_to_markdown(items)
                        if not inline:
                            inline = para.text.strip()
                        if inline:
                            texts.append(inline)
                    except Exception:
                        pass
                text = " ".join(texts)
                # Filter: skip empty text, pure digits (page numbers),
                # page-number patterns (e.g. "第  页", "第6页", "共 页"), and duplicates
                if (
                    text
                    and not text.strip().isdigit()
                    and not _RE_PAGE_ONLY.match(text.strip())
                    and text not in seen
                ):
                    seen.add(text)
                    results.append(text)
            except Exception:
                pass
        return results

    header_lines = []
    footer_lines = []
    seen_headers: set = set()
    seen_footers: set = set()

    # Check if the document uses different odd/even page headers/footers
    try:
        odd_even = doc.settings.odd_and_even_pages_header_footer
    except Exception:
        odd_even = False

    for section in doc.sections:
        try:
            # Collect headers
            hdrs = [section.header]
            if odd_even:
                try:
                    hdrs.append(section.even_page_header)
                except Exception:
                    pass
            try:
                if section.different_first_page_header_footer:
                    hdrs.append(section.first_page_header)
            except Exception:
                pass
            header_lines.extend(_collect_from_parts(hdrs, seen_headers))

            # Collect footers
            ftrs = [section.footer]
            if odd_even:
                try:
                    ftrs.append(section.even_page_footer)
                except Exception:
                    pass
            try:
                if section.different_first_page_header_footer:
                    ftrs.append(section.first_page_footer)
            except Exception:
                pass
            footer_lines.extend(_collect_from_parts(ftrs, seen_footers))
        except Exception:
            pass

    return header_lines, footer_lines


@default_registry.register
class DocxConverter(BaseConverter):
    supported_extensions = ["docx"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def convert_file(self, file_path: Path, **kwargs) -> ConvertResult:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "DOCX conversion requires python-docx: pip install paddleocr[doc2md]"
            )

        extract_drawings = kwargs.pop("extract_drawings", True)
        extract_headers_footers = kwargs.pop("extract_headers_footers", True)
        doc = Document(str(file_path))
        md_text, images = _convert_body(doc, extract_drawings=extract_drawings)

        if extract_headers_footers:
            header_lines, footer_lines = _extract_headers_footers(doc)
            parts = []
            if header_lines:
                parts.append("\n\n".join(header_lines))
                parts.append("")
            parts.append(md_text)
            if footer_lines:
                parts.append("")
                parts.append("---")
                parts.append("\n\n".join(footer_lines))
            md_text = "\n".join(parts)

        return ConvertResult(
            markdown=md_text,
            images=images,
            metadata={
                "format": "DOCX",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
